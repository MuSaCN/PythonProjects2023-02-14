# =============================================================================
# 13.1 演示文档
# =============================================================================

import docx
file = docx.Document()

file.add_paragraph('螃蟹在剥我的壳，笔记本在写我')
file.add_paragraph('漫天的我落在枫叶上雪花上')
file.add_paragraph('而你在想我')

file.save('E:\\三行情书.docx')
print('Word生成完毕！')